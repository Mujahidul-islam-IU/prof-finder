"""
ProfFinder — CV Parser Service
Extracts text from PDF (PyMuPDF) and DOCX (python-docx) files.
Uses GPT-4o-mini to structure the extracted text into a StudentProfile.
"""

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from openai import AsyncOpenAI
import json
from app.config import get_settings
from app.models.schemas import StudentProfile


# ── GPA Normalization Lookup ─────────────────────────────────
# Maps (country, scale) -> conversion factor to 4.0 scale
GPA_CONVERSION = {
    # Bangladesh: typically out of 4.0 or 5.0
    ("BD", "4.0"): lambda gpa: gpa,
    ("BD", "5.0"): lambda gpa: round(gpa * 4.0 / 5.0, 2),
    # Pakistan: out of 4.0
    ("PK", "4.0"): lambda gpa: gpa,
    # European ECTS: 1-10 scale
    ("EU", "10.0"): lambda gpa: round(gpa * 4.0 / 10.0, 2),
    # Australia: 7-point scale
    ("AU", "7.0"): lambda gpa: round(gpa * 4.0 / 7.0, 2),
    # UK: percentage -> 4.0
    ("UK", "100"): lambda gpa: round(min(gpa / 25.0, 4.0), 2),
    # US: already 4.0
    ("US", "4.0"): lambda gpa: gpa,
    # India: out of 10
    ("IN", "10.0"): lambda gpa: round(gpa * 4.0 / 10.0, 2),
    # Default: assume already 4.0
    ("DEFAULT", "4.0"): lambda gpa: gpa,
}


def normalize_gpa(gpa: float, country: str = "BD", scale: str = "4.0") -> float:
    """Normalize GPA to 4.0 scale using country-specific lookup."""
    key = (country.upper(), scale)
    converter = GPA_CONVERSION.get(key, GPA_CONVERSION[("DEFAULT", "4.0")])
    normalized = converter(gpa)
    return min(max(normalized, 0.0), 4.0)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file using PyMuPDF."""
    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    doc = Document(file_path)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


def extract_text_from_txt(file_path: str) -> str:
    """Extract all text from a TXT file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_cv_text(file_path: str) -> str:
    """Extract text from CV file (PDF or DOCX)."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF, DOCX, or TXT.")


PROFILE_EXTRACTION_PROMPT = """You are an expert CV/resume parser. Extract structured information from the following CV text.

Return a JSON object with these exact fields:
{
  "name": "Full name",
  "email": "Email address",
  "phone": "Phone number",
  "gpa_original": 3.5,  // numeric GPA, null if not found
  "gpa_scale": "4.0",   // the scale (4.0, 5.0, 10.0, 100), null if not found
  "degree_details": "BSc in Computer Science from XYZ University, 2023",
  "university": "University name",
  "ielts_score": 7.5,  // null if not found
  "gre_score": 320,    // null if not found
  "skills": ["Python", "Machine Learning", "NLP"],
  "research_interests": ["Natural Language Processing", "Computer Vision"],
  "publications": ["Paper title 1", "Paper title 2"],
  "thesis_summary": "Summary of thesis/research work if mentioned",
  "work_experience": "Brief summary of relevant work/research experience"
}

RULES:
- Extract ONLY information explicitly stated in the CV. Never infer or fabricate.
- If a field is not found, use null for optional fields or empty string/list.
- For GPA, identify both the value and the grading scale used.
- List ALL publications mentioned, including conference papers, journals, and preprints.
- For EACH publication, return ONLY the title of the paper. DO NOT include author names, journal names, years, or page numbers in the title string. (e.g. "Deep Learning for scRNA-seq" instead of "Islam et al., Deep Learning for scRNA-seq, Nature 2024").
- For research interests, combine explicitly stated interests with topics evident from publications/projects.

CV TEXT:
"""


async def parse_cv_with_llm(cv_text: str) -> StudentProfile:
    """Use GPT-4o-mini to extract structured profile from CV text."""
    settings = get_settings()
    client = AsyncOpenAI(api_key=settings.openai_api_key)

    response = await client.chat.completions.create(
        model=settings.openai_extraction_model,
        messages=[
            {"role": "system", "content": "You are a precise CV parser. Return valid JSON only."},
            {"role": "user", "content": PROFILE_EXTRACTION_PROMPT + cv_text}
        ],
        response_format={"type": "json_object"},
        temperature=0.0,
    )

    parsed = json.loads(response.choices[0].message.content)

    # Build profile
    profile = StudentProfile(
        name=parsed.get("name", ""),
        email=parsed.get("email", ""),
        phone=parsed.get("phone", ""),
        gpa_original=parsed.get("gpa_original"),
        gpa_scale=parsed.get("gpa_scale"),
        degree_details=parsed.get("degree_details", ""),
        university=parsed.get("university", ""),
        ielts_score=parsed.get("ielts_score"),
        gre_score=parsed.get("gre_score"),
        skills=parsed.get("skills", []),
        research_interests=parsed.get("research_interests", []),
        publications=parsed.get("publications", []),
        thesis_summary=parsed.get("thesis_summary", ""),
        work_experience=parsed.get("work_experience", ""),
        raw_text=cv_text,
    )

    # Normalize GPA if available
    if profile.gpa_original and profile.gpa_scale:
        profile.gpa_normalized = normalize_gpa(
            profile.gpa_original,
            country="BD",  # Default; can be parameterized later
            scale=profile.gpa_scale
        )

    return profile


async def parse_cv(file_path: str) -> StudentProfile:
    """Full CV parsing pipeline: extract text -> LLM parse -> normalize."""
    cv_text = extract_cv_text(file_path)
    profile = await parse_cv_with_llm(cv_text)
    return profile
